"""
DOCX Document Handler.

Provides utilities to extract text from a .docx file while preserving
structural metadata, and to write redacted content back to a new .docx
file at the run level so that bold, italic, and font attributes are retained.
"""

from __future__ import annotations

import copy
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

logger = logging.getLogger(__name__)


@dataclass
class TextSegment:
    """
    Represents a contiguous text block extracted from a docx element,
    along with its location so replacements can be written back precisely.
    """

    text: str
    paragraph_index: Optional[int] = None
    table_index: Optional[int] = None
    row_index: Optional[int] = None
    cell_index: Optional[int] = None
    is_table_cell: bool = False


class DocxReader:
    """
    Opens a .docx file and provides access to its text as a flat list of
    TextSegment objects, covering both body paragraphs and table cells.
    """

    def __init__(self, path: Path) -> None:
        self._path = path
        self._doc: Document = Document(str(path))

    @property
    def document(self) -> Document:
        """Return the underlying python-docx Document object."""
        return self._doc

    def extract_segments(self) -> List[TextSegment]:
        """
        Walk all paragraphs and table cells in document order and return
        a list of TextSegment objects, each carrying its location metadata.
        """
        segments: List[TextSegment] = []

        for para_idx, para in enumerate(self._doc.paragraphs):
            text = para.text
            if text.strip():
                segments.append(
                    TextSegment(
                        text=text,
                        paragraph_index=para_idx,
                        is_table_cell=False,
                    )
                )

        for tbl_idx, table in enumerate(self._doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        text = para.text
                        if text.strip():
                            segments.append(
                                TextSegment(
                                    text=text,
                                    table_index=tbl_idx,
                                    row_index=row_idx,
                                    cell_index=cell_idx,
                                    is_table_cell=True,
                                )
                            )

        logger.info("Extracted %d non-empty text segments from '%s'.", len(segments), self._path.name)
        return segments


def _replace_text_in_runs(runs: List[Run], original_text: str, replacement_text: str) -> bool:
    """
    Attempt to replace a substring across one or more consecutive runs while
    preserving the formatting of the first run in the matched span.
    Returns True if a replacement was made.
    """
    full_text = "".join(r.text for r in runs)
    if original_text not in full_text:
        return False

    new_full = full_text.replace(original_text, replacement_text, 1)

    if len(runs) == 1:
        runs[0].text = new_full
        return True

    first_run = runs[0]
    first_run.text = new_full
    for run in runs[1:]:
        run.text = ""
    return True


def _apply_replacement_to_paragraph(paragraph: Paragraph, original: str, replacement: str) -> None:
    """
    Apply a single PII replacement across all runs in a paragraph.
    Tries run-by-run first; falls back to a reconstructed single-run strategy
    that preserves the first run's character formatting.
    """
    para_text = paragraph.text
    if original not in para_text:
        return

    replaced = _replace_text_in_runs(paragraph.runs, original, replacement)

    if not replaced:
        new_text = para_text.replace(original, replacement, 1)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(new_text)


class DocxWriter:
    """
    Applies PII replacements to a python-docx Document in-place at the run level,
    then saves the result to a new output path.
    """

    def __init__(self, document: Document) -> None:
        self._doc = document

    def apply_replacements(self, replacements: List[tuple[str, str]]) -> None:
        """
        Apply a list of (original_text, replacement_text) tuples to every
        paragraph and table cell in the document. Optimized to filter candidates per element.
        """
        if not replacements:
            return

        replacements_sorted = sorted(replacements, key=lambda r: -len(r[0]))

        for paragraph in self._doc.paragraphs:
            para_text = paragraph.text
            if not para_text.strip():
                continue
            for original, replacement in replacements_sorted:
                if original in para_text:
                    _apply_replacement_to_paragraph(paragraph, original, replacement)
                    para_text = paragraph.text

        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        cell_text = paragraph.text
                        if not cell_text.strip():
                            continue
                        for original, replacement in replacements_sorted:
                            if original in cell_text:
                                _apply_replacement_to_paragraph(paragraph, original, replacement)
                                cell_text = paragraph.text

        logger.info("Applied %d replacements to document.", len(replacements_sorted))

    def save(self, output_path: Path) -> None:
        """Save the modified document to disk at the specified output path."""
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self._doc.save(str(output_path))
        logger.info("Redacted document saved to '%s'.", output_path)
